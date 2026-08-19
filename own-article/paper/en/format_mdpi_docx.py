"""Apply the official MDPI Agronomy Word styles to the generated manuscript.

The LaTeX sections remain the content source of truth.  ``make_docx.py`` first
converts them with the official Agronomy template as the reference document;
this script then maps the generated paragraphs to the template's named MDPI
styles, supplies the submission-style front matter, and writes a separate clean
DOCX so the previous manuscript is preserved.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "paper_en.docx"
OUTPUT = HERE / "paper_en_mdpi_simulation_short_captions.docx"


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    parent.remove(paragraph._p)


def paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    if style is not None:
        created.style = style
    if text:
        created.add_run(text)
    return created


def prepend_run(paragraph: Paragraph, text: str, *, bold: bool = False) -> None:
    run = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        run.append(rpr)
    node = OxmlElement("w:t")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    ppr = paragraph._p.pPr
    insert_at = 1 if ppr is not None else 0
    paragraph._p.insert(insert_at, run)


def remove_text_prefix(paragraph: Paragraph, prefix: str) -> None:
    """Remove a plain-text prefix without discarding later run formatting."""
    remaining = len(prefix)
    for run in paragraph.runs:
        if remaining == 0:
            break
        if len(run.text) <= remaining:
            remaining -= len(run.text)
            run.text = ""
        else:
            run.text = run.text[remaining:]
            remaining = 0
    if remaining:
        raise RuntimeError(f"Could not remove paragraph prefix: {prefix!r}")


def add_author_line(paragraph: Paragraph) -> None:
    clear_paragraph(paragraph)
    r = paragraph.add_run("[[AUTHOR 1 - FULL NAME REQUIRED]]")
    r.bold = True
    r = paragraph.add_run("1,*")
    r.font.superscript = True
    paragraph.add_run("; ")
    r = paragraph.add_run("[[AUTHOR 2 - FULL NAME REQUIRED]]")
    r.bold = True
    r = paragraph.add_run("2")
    r.font.superscript = True
    paragraph.add_run("; [[ADD OR DELETE AUTHORS AS REQUIRED]]")


def set_labeled_placeholder(paragraph: Paragraph, prefix: str, remainder: str) -> None:
    clear_paragraph(paragraph)
    r = paragraph.add_run(prefix)
    r.bold = True
    paragraph.add_run(remainder)


def set_front_matter(doc: Document) -> None:
    original = list(doc.paragraphs)
    if len(original) < 7:
        raise RuntimeError("Unexpected generated front matter")

    title, author1, author2, author3, abstract_title, abstract, keywords = original[:7]

    article = title.insert_paragraph_before("Article")
    article.style = doc.styles["MDPI_1.1_article_type"]

    title.style = doc.styles["MDPI_1.2_title"]

    author1.style = doc.styles["MDPI_1.3_authornames"]
    add_author_line(author1)

    author2.style = doc.styles["MDPI_1.6_affiliation"]
    set_labeled_placeholder(
        author2,
        "1 ",
        "[[AFFILIATION 1 REQUIRED]]; [[AUTHOR 1 E-MAIL REQUIRED]]",
    )

    author3.style = doc.styles["MDPI_1.6_affiliation"]
    set_labeled_placeholder(
        author3,
        "2 ",
        "[[AFFILIATION 2 REQUIRED]]; [[AUTHOR 2 E-MAIL REQUIRED]]",
    )

    correspondence = paragraph_after(author3, style=doc.styles["MDPI_1.6_affiliation"])
    set_labeled_placeholder(
        correspondence,
        "* Correspondence: ",
        "[[CORRESPONDING AUTHOR E-MAIL REQUIRED]]",
    )

    remove_paragraph(abstract_title)
    abstract.style = doc.styles["MDPI_1.7_abstract"]
    prepend_run(abstract, "Abstract: ", bold=True)

    keywords.style = doc.styles["MDPI_1.8_keywords"]
    line = paragraph_after(keywords, style=doc.styles["MDPI_1.9_line"])
    line.alignment = WD_ALIGN_PARAGRAPH.LEFT


def style_body(doc: Document) -> None:
    h1_no = 0
    h2_no = 0
    table_no = 0
    figure_no = 0
    in_references = False

    for paragraph in doc.paragraphs:
        ppr = paragraph._p.pPr
        style_id = (
            ppr.pStyle.val
            if ppr is not None and ppr.pStyle is not None
            else ""
        )
        text = paragraph.text.strip()

        if style_id == "Heading1":
            if text == "References":
                paragraph.style = doc.styles["MDPI_2.1_heading1"]
                in_references = True
            else:
                h1_no += 1
                h2_no = 0
                paragraph.style = doc.styles["MDPI_2.1_heading1"]
                prepend_run(paragraph, f"{h1_no}. ")
            continue

        if in_references:
            numbered_prefix = re.match(r"^\d+\.\s+", text)
            if numbered_prefix:
                remove_text_prefix(paragraph, numbered_prefix.group(0))
            paragraph.style = doc.styles["MDPI_8.1_references"]
            continue

        if style_id == "Heading2":
            h2_no += 1
            paragraph.style = doc.styles["MDPI_2.2_heading2"]
            prepend_run(paragraph, f"{h1_no}.{h2_no}. ")
            continue

        if style_id == "Heading3":
            paragraph.style = doc.styles["MDPI_2.3_heading3"]
            continue

        if style_id == "TableCaption":
            table_no += 1
            paragraph.style = doc.styles["MDPI_4.1_table_caption"]
            prepend_run(paragraph, f"Table {table_no}. ", bold=True)
            continue

        if style_id == "ImageCaption":
            paragraph.style = doc.styles["MDPI_5.1_figure_caption"]
            if text.lower().startswith("graphical abstract"):
                prepend_run(paragraph, "Graphical Abstract. ", bold=True)
                # Drop the duplicate words already at the start while retaining
                # the caption's inline mathematics and emphasis.
                for run in paragraph.runs[1:]:
                    if run.text.lower().startswith("graphical abstract. "):
                        run.text = run.text[len("Graphical abstract. ") :]
                        break
            else:
                figure_no += 1
                prepend_run(paragraph, f"Figure {figure_no}. ", bold=True)
            continue

        if style_id == "CaptionedFigure":
            paragraph.style = doc.styles["MDPI_5.2_figure"]
            continue

        if re.fullmatch(r"\(\d+\)", text):
            paragraph.style = doc.styles["MDPI_3.a_equation_number"]
            continue

        has_display_math = bool(paragraph._p.xpath("./m:oMath | ./m:oMathPara"))
        if has_display_math and not text:
            paragraph.style = doc.styles["MDPI_3.9_equation"]
            continue

        if style_id in {"BodyText", "FirstParagraph"}:
            paragraph.style = doc.styles["MDPI_3.1_text"]


def add_and_style_back_matter(doc: Document) -> None:
    labels = {
        "Author Contributions:",
        "Funding:",
        "Institutional Review Board Statement:",
        "Informed Consent Statement:",
        "Data Availability Statement:",
        "Conflicts of Interest:",
    }
    paragraphs = list(doc.paragraphs)
    author_contrib = next(p for p in paragraphs if p.text.startswith("Author Contributions:"))
    supplementary = author_contrib.insert_paragraph_before()
    supplementary.style = doc.styles["MDPI_6.2_back_matter"]
    set_labeled_placeholder(supplementary, "Supplementary Materials: ", "Not applicable.")

    conflicts = next(p for p in doc.paragraphs if p.text.startswith("Conflicts of Interest:"))
    acknowledgments = conflicts.insert_paragraph_before()
    acknowledgments.style = doc.styles["MDPI_6.2_back_matter"]
    set_labeled_placeholder(
        acknowledgments,
        "Acknowledgments: ",
        "[[ACKNOWLEDGMENTS REQUIRED; OTHERWISE STATE 'Not applicable.']]",
    )

    for paragraph in doc.paragraphs:
        if any(paragraph.text.startswith(label) for label in labels):
            paragraph.style = doc.styles["MDPI_6.2_back_matter"]


def style_tables(doc: Document) -> None:
    for table in doc.tables:
        table.style = doc.styles["MDPI_4.1_three_line_table"]
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() == "2-3(lr)4-5 Library":
                        clear_paragraph(paragraph)
                        paragraph.add_run("Library")
                    paragraph.style = doc.styles["MDPI_4.2_table_body"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def validate(doc: Document) -> None:
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "in-silico title": "an in silico study" in text,
        "simulation scope in abstract": "comparative simulation evidence" in text,
        "physical-validation limitation": "There is no physical-greenhouse validation" in text,
        "five numbered sections": all(f"{i}. " in text for i in range(1, 6)),
        "sixteen table captions": sum(
            p.style.style_id == "MDPI41tablecaption" for p in doc.paragraphs
        ) == 16,
        "five numbered figures": sum(
            p.text.startswith("Figure ") for p in doc.paragraphs
        ) == 5,
        "graphical abstract unnumbered": any(
            p.text.startswith("Graphical Abstract.") for p in doc.paragraphs
        ),
        "forty-eight references": sum(
            p.style.style_id == "MDPI81references" for p in doc.paragraphs
        ) == 48,
        "back matter": sum(
            p.style.style_id == "MDPI62backmatter" for p in doc.paragraphs
        ) == 8,
    }
    for label, passed in checks.items():
        print(("OK   " if passed else "FAIL ") + label)
    if not all(checks.values()):
        raise RuntimeError("MDPI formatting validation failed")


def main() -> None:
    doc = Document(SOURCE)
    set_front_matter(doc)
    style_body(doc)
    add_and_style_back_matter(doc)
    style_tables(doc)
    validate(doc)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    main()
